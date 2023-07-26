import os
import pandas as pd
import pdf2image
from azure.cognitiveservices.vision.computervision import ComputerVisionClient
from msrest.authentication import CognitiveServicesCredentials
from docx import Document
from io import BytesIO
import time
import logging


def getcreds():
    with open("creds_cv.txt", "r") as c:
        creds = c.readlines()
    return creds[0].strip(), creds[1].strip()


class DocClient:
    def __init__(self, endpoint, key):
        self.client = ComputerVisionClient(endpoint, CognitiveServicesCredentials(key))

    def close(self):
        self.client.close()

    def extract_content(self, result):
        contents = {
            "page": [],
            "content": [],
            "confidence": [],
        }

        for read_result in result.analyze_result.read_results:
            lines = read_result.lines
            lines.sort(key=lambda line: line.bounding_box[1])

            for line in lines:
                contents["page"].append(read_result.page)
                contents["content"].append(" ".join([word.text for word in line.words]))
                contents["confidence"].append(
                    sum([word.confidence for word in line.words]) / len(line.words)
                )

        return pd.DataFrame(contents)

    def pdf2df(self, pdf_path):
        with open(pdf_path, "rb") as file:
            pdf_data = file.read()

            num_pages = pdf2image.pdfinfo_from_bytes(pdf_data)["Pages"]

            results = []
            for i in range(num_pages):
                try:
                    image = pdf2image.convert_from_bytes(
                        pdf_data, dpi=500, first_page=i + 1, last_page=i + 1
                    )[0]

                    img_byte_arr = BytesIO()
                    image.save(img_byte_arr, format="PNG")

                    img_byte_arr.seek(0)
                    ocr_result = self.client.read_in_stream(img_byte_arr, raw=True)
                    operation_id = ocr_result.headers["Operation-Location"].split("/")[
                        -1
                    ]

                    while True:
                        result = self.client.get_read_result(operation_id)

                        if result.status.lower() not in ["notstarted", "running"]:
                            break

                        time.sleep(1)

                    if result.status.lower() == "failed":
                        logging.error(f"OCR failed for page {i+1} of file {pdf_path}")
                        continue

                    df_results = self.extract_content(result)
                    results.append(df_results)
                except azure.core.exceptions.HttpResponseError as e:
                    logging.error(
                        f"Error processing page {i+1} of file {pdf_path}: {e}"
                    )
                    continue

            combined_results = pd.concat(results)

            return combined_results

    def process(self, pdf_path):
        outname = os.path.basename(pdf_path).replace(".pdf", "")
        outstring = os.path.join(f_path, "{}.docx".format(outname))
        outpath = os.path.abspath(outstring)
        if os.path.exists(outpath):
            logging.info(f"skipping {outpath}, file already exists")
            return outpath

        logging.info(f"sending document {outname}")
        results = self.pdf2df(pdf_path)
        logging.info(f"writing to {outpath}")

        doc = Document()
        for page, group in results.groupby("page"):
            doc.add_paragraph("\n".join(group["content"].tolist()))
            doc.add_paragraph(
                "\n" + "-" * 10 + f" End of Page {page} " + "-" * 10 + "\n"
            )

        doc.save(outpath)

        return outpath
